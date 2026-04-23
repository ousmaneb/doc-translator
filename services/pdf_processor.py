"""
PDF processing: extract text with positions (pdfminer.six), then overlay
translated text onto the original PDF (reportlab + pypdf).
"""
import io
import re

try:
    from pdfminer.high_level import extract_pages
    from pdfminer.layout import LTTextBox, LTTextLine, LTChar
    PDFMINER_OK = True
except ImportError:
    PDFMINER_OK = False

try:
    from pypdf import PdfReader, PdfWriter
    PYPDF_OK = True
except ImportError:
    PYPDF_OK = False

try:
    from reportlab.pdfgen import canvas as rl_canvas
    REPORTLAB_OK = True
except ImportError:
    REPORTLAB_OK = False

try:
    from docx import Document as DocxDoc
    from docx.shared import Pt
    DOCX_OK = True
except ImportError:
    DOCX_OK = False


# ── extraction ────────────────────────────────────────────────────────────────

def extract_pdf_pages(buffer: bytes) -> list[dict]:
    if not PDFMINER_OK:
        raise ImportError("Install: pip install pdfminer.six")

    pdf_io = io.BytesIO(buffer)
    pages: list[dict] = []

    for page_layout in extract_pages(pdf_io):
        width = float(page_layout.width)
        height = float(page_layout.height)
        raw: list[dict] = []

        for element in page_layout:
            if not isinstance(element, LTTextBox):
                continue
            for line in element:
                if not isinstance(line, LTTextLine):
                    continue
                text = line.get_text().rstrip("\n").strip()
                if not text:
                    continue
                if re.search(r"<[a-zA-Z]", text) or "</" in text:
                    continue
                fs = 10.0
                for char in line:
                    if isinstance(char, LTChar):
                        fs = float(char.size)
                        break
                raw.append({
                    "str": text,
                    "x": float(line.x0),
                    "y": float(line.y0),
                    "fs": fs,
                    "w": float(line.x1 - line.x0),
                })

        if not raw:
            pages.append({"lines": [], "width": width, "height": height, "median_font_size": 10.0})
            continue

        sorted_fs = sorted(r["fs"] for r in raw)
        med_fs = sorted_fs[len(sorted_fs) // 2]

        raw.sort(key=lambda r: (-r["y"], r["x"]))

        y_thresh = med_fs * 0.65
        groups: list[list[dict]] = []
        for item in raw:
            if groups and abs(item["y"] - groups[-1][0]["y"]) <= y_thresh:
                groups[-1].append(item)
            else:
                groups.append([item])

        lines: list[dict] = []
        for group in groups:
            group.sort(key=lambda r: r["x"])
            avg_fs = sum(r["fs"] for r in group) / len(group)
            x_span = (group[-1]["x"] + group[-1]["w"]) - group[0]["x"]
            is_table_row = len(group) > 1 and x_span > width * 0.2
            is_heading = (avg_fs / med_fs >= 1.3) if med_fs > 0 else False
            cells = [{"text": r["str"], "x": r["x"], "w": r["w"], "translated": ""} for r in group]
            lines.append({
                "cells": cells,
                "y": group[0]["y"],
                "font_size": avg_fs,
                "is_heading": is_heading,
                "is_table_row": is_table_row,
            })

        lines = _normalize_table_columns(lines, width)
        pages.append({"lines": lines, "width": width, "height": height, "median_font_size": med_fs})

    return pages


def _cluster_columns(rows: list[dict], page_width: float) -> list[float]:
    all_x = sorted(cell["x"] for row in rows for cell in row["cells"])
    cluster_pt = max(15.0, page_width * 0.04)
    cols: list[float] = []
    for x in all_x:
        if not cols or x - cols[-1] > cluster_pt:
            cols.append(x)
    return cols


def _normalize_table_columns(lines: list[dict], page_width: float) -> list[dict]:
    out: list[dict] = []
    i = 0
    while i < len(lines):
        if not lines[i]["is_table_row"]:
            out.append(lines[i])
            i += 1
            continue

        section: list[dict] = []
        while i < len(lines) and lines[i]["is_table_row"]:
            section.append(lines[i])
            i += 1

        cols = _cluster_columns(section, page_width)
        if len(cols) < 2:
            out.extend(section)
            continue

        for row in section:
            col_map: dict[float, dict] = {}
            for cell in row["cells"]:
                nearest = min(cols, key=lambda c: abs(cell["x"] - c))
                if nearest in col_map:
                    col_map[nearest]["text"] += " " + cell["text"]
                    col_map[nearest]["w"] += cell["w"]
                else:
                    col_map[nearest] = {"text": cell["text"], "w": cell["w"]}

            cells = [
                {"text": col_map[c]["text"].strip(), "x": c, "w": col_map[c]["w"], "translated": ""}
                for c in cols if c in col_map
            ]
            if cells:
                out.append({**row, "cells": cells})

    return out


# ── PDF overlay ───────────────────────────────────────────────────────────────

def _make_overlay(page_data: dict, pw: float, ph: float) -> io.BytesIO:
    if not REPORTLAB_OK:
        raise ImportError("Install: pip install reportlab")

    packet = io.BytesIO()
    c = rl_canvas.Canvas(packet, pagesize=(pw, ph))

    for line in page_data["lines"]:
        cells = line["cells"]
        if not cells:
            continue

        fs = max(6.5, min(36.0, line["font_size"]))
        orig_fs = max(6.5, line["font_size"])
        bold = line["is_heading"]
        font_name = "Helvetica-Bold" if bold else "Helvetica"

        rect_y = line["y"] - orig_fs * 0.3
        rect_h = orig_fs * 1.4
        if rect_y < -orig_fs or rect_y > ph + orig_fs:
            continue

        if line["is_table_row"] and len(cells) > 1:
            for ci, cell in enumerate(cells):
                cell_text = cell.get("translated") or cell["text"]
                cell_x = cell["x"]
                col_end = cells[ci + 1]["x"] - 3 if ci < len(cells) - 1 else pw - 5
                col_w = max(col_end - cell_x, 10.0)
                if cell_x >= pw - 5:
                    continue

                orig_w = cell["w"]
                try:
                    trans_w = c.stringWidth(cell_text, font_name, fs)
                except Exception:
                    trans_w = len(cell_text) * fs * 0.55

                rect_w = min(max(orig_w, trans_w) + 4, col_w)
                c.setFillColorRGB(1, 1, 1)
                c.rect(cell_x - 2, rect_y, rect_w + 2, rect_h, fill=1, stroke=0)

                draw_fs = max(6.0, fs * col_w / trans_w) if trans_w > col_w else fs
                c.setFillColorRGB(0, 0, 0)
                c.setFont(font_name, draw_fs)
                try:
                    c.drawString(cell_x, line["y"], cell_text)
                except Exception:
                    pass
        else:
            text = " ".join(cell.get("translated") or cell["text"] for cell in cells)
            x = cells[0]["x"]
            orig_w = sum(cell["w"] for cell in cells)
            try:
                trans_w = c.stringWidth(text, font_name, fs)
            except Exception:
                trans_w = len(text) * fs * 0.55

            avail_w = max(pw - x - 4, 50.0)
            rect_w = min(max(orig_w, trans_w) + 6, avail_w + 2)
            c.setFillColorRGB(1, 1, 1)
            c.rect(x - 2, rect_y, rect_w + 2, rect_h, fill=1, stroke=0)

            c.setFillColorRGB(0, 0, 0)
            c.setFont(font_name, fs)

            if trans_w <= avail_w:
                try:
                    c.drawString(x, line["y"], text)
                except Exception:
                    pass
            else:
                words = text.split()
                cur = ""
                y = line["y"]
                lh = fs * 1.25
                for word in words:
                    test = f"{cur} {word}" if cur else word
                    try:
                        tw = c.stringWidth(test, font_name, fs)
                    except Exception:
                        tw = len(test) * fs * 0.55
                    if tw > avail_w and cur:
                        try:
                            c.drawString(x, y, cur)
                        except Exception:
                            pass
                        y -= lh
                        cur = word
                    else:
                        cur = test
                if cur:
                    try:
                        c.drawString(x, y, cur)
                    except Exception:
                        pass

    c.save()
    packet.seek(0)
    return packet


def build_translated_pdf(pages: list[dict], original_buffer: bytes, output_path: str) -> None:
    if not PYPDF_OK or not REPORTLAB_OK:
        raise ImportError("Install: pip install pypdf reportlab")

    reader = PdfReader(io.BytesIO(original_buffer))
    writer = PdfWriter()

    for i in range(min(len(pages), len(reader.pages))):
        orig_page = reader.pages[i]
        pw = float(orig_page.mediabox.width)
        ph = float(orig_page.mediabox.height)

        if pages[i]["lines"]:
            overlay = _make_overlay(pages[i], pw, ph)
            overlay_reader = PdfReader(overlay)
            orig_page.merge_page(overlay_reader.pages[0])

        writer.add_page(orig_page)

    with open(output_path, "wb") as f:
        writer.write(f)


# ── build simple PDF from plain text ─────────────────────────────────────────

def build_simple_pdf(text: str, output_path: str) -> None:
    if not REPORTLAB_OK:
        raise ImportError("Install: pip install reportlab")
    from reportlab.lib.pagesizes import A4

    pw, ph = A4
    margin = 50.0
    fs = 11.0
    lh = fs * 1.55
    max_w = pw - 2 * margin

    c = rl_canvas.Canvas(output_path, pagesize=A4)
    c.setFont("Helvetica", fs)
    c.setFillColorRGB(0, 0, 0)
    y = ph - margin - fs

    def new_page():
        nonlocal y
        c.showPage()
        c.setFont("Helvetica", fs)
        c.setFillColorRGB(0, 0, 0)
        y = ph - margin - fs

    for para in text.split("\n"):
        if not para.strip():
            y -= lh * 0.5
            if y < margin:
                new_page()
            continue

        words = para.split()
        cur = ""
        for word in words:
            test = f"{cur} {word}" if cur else word
            try:
                tw = c.stringWidth(test, "Helvetica", fs)
            except Exception:
                tw = len(test) * fs * 0.55
            if tw > max_w and cur:
                if y < margin + lh:
                    new_page()
                try:
                    c.drawString(margin, y, cur)
                except Exception:
                    pass
                y -= lh
                cur = word
            else:
                cur = test

        if cur:
            if y < margin + lh:
                new_page()
            try:
                c.drawString(margin, y, cur)
            except Exception:
                pass
            y -= lh

        y -= lh * 0.3
        if y < margin:
            new_page()

    c.save()


# ── build structured DOCX from PDF pages ─────────────────────────────────────

def build_docx_from_pages(pages: list[dict], output_path: str) -> None:
    if not DOCX_OK:
        raise ImportError("Install: pip install python-docx")

    doc = DocxDoc()

    for page_idx, page in enumerate(pages):
        med_fs = page["median_font_size"] or 10.0
        i = 0
        while i < len(page["lines"]):
            line = page["lines"][i]

            if line["is_table_row"]:
                table_lines: list[dict] = []
                while i < len(page["lines"]) and page["lines"][i]["is_table_row"]:
                    table_lines.append(page["lines"][i])
                    i += 1

                max_cols = max(len(tl["cells"]) for tl in table_lines)
                if max_cols <= 0:
                    continue

                table = doc.add_table(rows=len(table_lines), cols=max_cols)
                table.style = "Table Grid"

                for ri, tl in enumerate(table_lines):
                    cells = list(tl["cells"])
                    while len(cells) < max_cols:
                        cells.append({"text": "", "x": 0, "w": 0, "translated": ""})
                    for ci, cell in enumerate(cells):
                        cell_text = cell.get("translated") or cell["text"]
                        tc = table.cell(ri, ci)
                        tc.text = cell_text
                        if tc.paragraphs and tc.paragraphs[0].runs:
                            run = tc.paragraphs[0].runs[0]
                            run.font.size = Pt(max(6.5, tl["font_size"]))
                            run.font.bold = tl["is_heading"]

                doc.add_paragraph()
            else:
                text = " ".join(cell.get("translated") or cell["text"] for cell in line["cells"]).strip()
                if text:
                    ratio = line["font_size"] / med_fs if med_fs else 1.0
                    if ratio >= 1.8:
                        doc.add_heading(text, level=1)
                    elif ratio >= 1.3:
                        doc.add_heading(text, level=2)
                    else:
                        para = doc.add_paragraph()
                        run = para.add_run(text)
                        run.font.size = Pt(max(6.5, min(72.0, line["font_size"])))
                        run.font.bold = line["is_heading"]
                i += 1

        if page_idx < len(pages) - 1:
            doc.add_page_break()

    doc.save(output_path)


# ── build simple DOCX from plain text ────────────────────────────────────────

def build_simple_docx(text: str, output_path: str) -> None:
    if not DOCX_OK:
        raise ImportError("Install: pip install python-docx")
    doc = DocxDoc()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(output_path)
